"""
Task 1 - Method Statement Automation

- Use the BOQ Extract and Project Details Sheet to generate a draft Method Statement for blockwork & plastering.
- Populate at least the following sections: Scope, References, Materials, Work Procedure.
"""

import json
import os
import io
from io import BytesIO
from task1.BOQExtract import BOQExtractParser, MethodStatement
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.encoders import jsonable_encoder
from fastapi import FastAPI, File, UploadFile, HTTPException

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utilities.exceptions import UnsupportedFileTypeError

app = FastAPI()

def _create_ms_document(boq_data: MethodStatement) -> io.BytesIO:
    """
    Create a Method Statement document from the provided data.
    """
    try:
        document = Document()
        
        # Title
        title = document.add_heading('Method Statement for Blockwork & Plastering', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        document.add_paragraph()  # Add a blank line

        def add_section(heading, content):
            document.add_heading(heading, level=2)
            if content is None:
                content = "N/A"
            
            if "\n" in content:
                items = [item.strip() for item in content.split("\n") if item.strip()]
                for item in items:
                    p = document.add_paragraph(item, style='List Bullet')
            else:
                document.add_paragraph(content)
            document.add_paragraph()  # Add a blank line

        # Populate sections
        if boq_data.MethodStatement:
            data = boq_data.MethodStatement[0]
            add_section("Scope", data.Scope or "N/A")
            add_section("References", data.References or "N/A")
            add_section("Materials", data.Materials or "N/A")
            add_section("Work Procedure", data.WorkProcedure or "N/A")

        # add general notes
        add_section("General Notes", boq_data.general_notes or "N/A")
        add_section("Follow-up Instructions", boq_data.follow_up_instructions or "N/A")

        # Save to BytesIO 
        doc_io = io.BytesIO()
        document.save(doc_io)
        doc_io.seek(0) # Reset pointer to the beginning
        return doc_io

    except Exception as e:
        print(f"Error creating document: {e}")
        raise e
    

@app.post("/task1")
async def process_boq_document(file: UploadFile = File(...)):
    try:
        boq_parser = BOQExtractParser()
        boq_data = await boq_parser.parse_boq_document(file.file)
        
        # return JSONResponse(content=jsonable_encoder(boq_data))
        document_stream = _create_ms_document(boq_data)

        # save to file
        output_path = "Method_Statement_Output.docx"
        with open(output_path, "wb") as f:
            f.write(document_stream.getvalue())
        
        document_stream.seek(0)

        return StreamingResponse(
            document_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Method_Statement.docx"}
        )
    
    except UnsupportedFileTypeError as e:
        raise HTTPException(status_code=400, detail="Unsupported file type.")
    
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)